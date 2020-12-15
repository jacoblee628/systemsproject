import pathlib
from pathlib import Path

import docx
import openpyxl as pyxl
import pandas as pd
from docx import Document
from openpyxl import Workbook

import re


def read_trace(file_path, matrix_type, return_df=True):
    """Loads trace matrix .xlsx file and extracts data

    Args:
        file_path (String): path to the trace file
        matrix_type (String): "CO" or "PSC"
        return_df (bool, optional): If true (default), returns a pandas dataframe. Else, returns list of lists.

    Returns:
        list or pd.DataFrame: the trace data
    """
    print(f"Loading {matrix_type} from: {file_path}")

    # Some asserts to make sure inputs are valid
    matrix_type = matrix_type.upper()
    assert matrix_type in ["CO", "PSC"], f"Trace matrix type must be either 'CO' or 'PSC'\nCurrent type: {matrix_type}"

    # Load in the actual excel file
    wb = pyxl.load_workbook(file_path)

    # Get the appropriate matrix from the workbook
    if matrix_type == "CO":
        ws = wb['CO Trace Matrix']
    else:
        ws = wb['PSC Trace Matrix']

    # get title of workbook
    title = ws['A1'].value
    print(f"Loading in worksheet titled: {title}")

    # extract headers for the table
    headers = [cell.value for cell in ws['A3':'E3'][0]]

    # load in the data row by row
    data = []
    for row in ws.iter_rows(min_row=4, max_col=5, values_only=True):
        row_data = [cell for cell in row]

        # Stop if all the row data is none
        # apparently the blank rows in the bottom of the worksheet are populated...
        if all([cell == None for cell in row_data]):
            break
        data.append(row_data)

    # Returns either pandas dataframe or a list of lists
    if return_df:
        return pd.DataFrame(data, columns=headers)
    else:
        return data

def read_as_run_tests(file_path, return_df=True):
    """Loads in the test statuses from the manual as run tests document. Document must be in .docx format.

    Args:
        file_path (String): path to the file (.docx)
        return_df (bool, optional): If true, returns pandas dataframe. Else dict. Defaults to True.

    Returns:
        pd.DataFrame or dict: Test names and corresponding statuses
    """
    # Open the file and read with python-docx package
    with open(file_path, 'rb') as f:
        document = Document(f)
        
    # Get the test names
    # (Need to iterate separately because theyre technically separate "paragraphs")
    test_names = []
    run_ids = []
    # For each "paragraph" in the document,
    for i, paragraph in enumerate(document.paragraphs):
        # Find lines that have "Run ID" in them
        if "Run ID:" in paragraph.text:
            # The line before that has the test name
            test_names.append(document.paragraphs[i-1].text)
            # Also store the run_id (for error log)
            run_id = int(paragraph.text.replace("Run ID:", ""))
            run_ids.append(run_id)

    # Get the table info
    data = []

    # For each table in the document:
    entry_idx = 0
    for table in document.tables:
        # If the first few cells belong to a test, it's probably a test description table
        if table.cell(0,0).text == "Status:" and table.cell(1,2).text == "1. Product":
            # Extract the test info from table
            data.append({
                "Test Name": test_names[entry_idx],
                "Run ID": run_ids[entry_idx],
                "Test Status": table.cell(0,1).text,
                "Release":table.cell(2,1).text,
                "Application": table.cell(1,3).text,
            })

            entry_idx += 1

    # make a dataframe for outputting
    df = pd.DataFrame(data)

    # Also add the V&V Test Report info by extracting it from file name
    file_name = Path(file_path).stem
    v_v = re.findall("ER([0-9]+ v[0-9]+|[0-9]+v[0-9]+)", file_name)[0].replace("ER", "")
    df["V&V Test Report"] = v_v

    # Output as either pandas dataframe or dict, depending on return_df setting.
    if return_df:
        return df
    else:
        return df.to_dict('records')


def read_msgateway_results(file_path, return_df=True):
    """Loads in the test statuses from the automatic MSGateway tests document. Document must be in .docx format.

    Args:
        file_path (String): path to the file
        return_df (bool, optional): If true, returns pandas dataframe. Else dict. Defaults to True.

    Returns:
        pd.DataFrame or dict: Test names and corresponding statuses
    """
    # Read the document
    with open(file_path, 'rb') as f:
        document = Document(f)

    # Get the names of the tests
    test_names = []
    for p in document.paragraphs:
        if "Test:" in p.text and "\t" not in p.text:
            test_names.append(p.text.replace("Test: ", ""))

    statuses = []
    for table in document.tables:
        if "Execution Status" in table.cell(0,0).text:
            statuses.append(table.cell(1,0).text)

    # Make sure same number of test names and statuses
    assert len(test_names) == len(statuses), f"Error: Couldn't parse same number of test names and test statuses.\n# test names: {len(test_names)}, # statuses: {len(statuses)}"

    # Output as either pandas dataframe or dict, depending on return_df setting.
    data = {"test_name":test_names, "status":statuses}

    if return_df:
        return pd.DataFrame(data)
    else:
        return data


def read_performance_test_results_by_tc(file_path, return_df=True):
    """Loads in the test statuses from the automatic PerformanceTestResultsByTC excel sheet.
    Document must be in .xlsx format.

    NOTE: Noticed it's manually filled out, so this may break.
          You can change the column names, and add/remove columns,
          but don't leave any blank rows between data you want to have processed.

    Args:
        file_path (String): path to the file
        return_df (bool, optional): If true, returns pandas dataframe. Else dict. Defaults to True.

    Returns:
        pd.DataFrame or dict: Test names and corresponding statuses
    """
    # Load workbook
    wb = pyxl.load_workbook(file_path)

    # Read table values (includes header at row 1)
    contents = []
    for row in wb.worksheets[0].values:
        # Break if all the values in the row are empty
        if all([value is None for value in row]):
            break
        contents.append(row)

    # First row is column headers, onwards is data
    df = pd.DataFrame(contents[1:], columns=contents[0])
    
    # Format to be a bit more similar to other datasets    
    df = df.rename({"PASS/FAIL": "Test Status"}, axis=1)
    df["Test Status"] = df["Test Status"].replace("PASS", "Passed").replace("FAIL", "Failed")

    # Return pandas dataframe if specified, otherwise return a list of dicts
    if return_df:
        return df
    else:
        return df.to_dict('records')


def read_rest_api_tests(folder_path, return_df=True):
    """Reads in all the rest api automatic test .txt files.
    
    Will recursively go through the folders and read in .txt files.
    Make sure everything is unzipped first; in future can add functionality to
    unzip automatically if needed.

    Args:
        folder_path (String or pathlib.Path): path to the main RestApiTests folder.
        return_df (bool, optional): If true, returns pandas dataframe. Else dict. Defaults to True.

    Returns:
        pd.DataFrame or dict: Test names and corresponding statuses
    """
    print(f"Loading api test files from {folder_path}")
    if isinstance(folder_path, str):
        folder_path = Path(folder_path)

    # Read all .txt files recursively in the folders
    file_list = [file_name for file_name in folder_path.rglob("*.txt") if "__MACOSX" not in str(file_name)]
    
    # Load in the data from each file, add to a single list
    data = []
    for file_name in file_list:
        data.extend(_read_group_by_method_txt(file_name, return_df=False))

    # make a dataframe for outputting
    df = pd.DataFrame(data)
    
    # Output as either pandas dataframe or dict, depending on return_df setting.
    if return_df:
        return df
    else:
        return df.to_dict('records')


def _read_group_by_method_txt(file_path, return_df=True):
    """Reads in a single rest api automatic test .txt file.

    Mostly just for the `read_rest_api_tests()` method. Usually won't need to call yourself

    Args:
        file_path (String): path to the .txt file
        return_df (bool, optional): If true, returns pandas dataframe. Else dict. Defaults to True.

    Returns:
        pd.DataFrame or dict: Test names and corresponding statuses
    """
    # If receive str, convert to pathlib object
    if isinstance(file_path, str):
        file_path = Path(file_path)

    # Read in file using pathlib method, split on new line symbols
    lines = file_path.read_text().split("\n")

    # Remove blank lines
    lines = [line for line in lines if len(line) > 0]

    data = []
    for line in lines:
        if len(line) < 2:
            continue # skip any lines that are too short; likely empty
        
        # Remove the "com.philips.sapphire.systemintegrationtests." part
        line = line.replace("com.philips.sapphire.systemintegrationtests.", "")
        
        # split on the | char; should only be one of them
        line_data = line.split("|")

        # For some reason if theres multiple of that char, only consider the last one
        if len(line_data) > 2:
            print(f"found line with multiple '|' chars: {line}")
            line_data = "".join(line_data[:-1]) + [line_data[-1]]
        
        # Get relevant data for other columns
        base_folder_idx = [i for i, s in enumerate(file_path.parts) if "RestApiTests" in str(s)][0]
        
        # Also add the V&V Test Report info by extracting it from file name
        er_folder_name = file_path.parts[base_folder_idx - 2]
        v_v = re.findall("ER([0-9]+ v[0-9]+|[0-9]+v[0-9]+)", er_folder_name)[0].replace("ER", "")
        
        # If there's no space in the file name (like ER2228014v53), add one
        v_idx = v_v.find("v")
        if v_v[v_idx - 1] != " ":
            v_v = v_v[:v_idx] + " " + v_v[v_idx:]
        
        # Store data in list
        data.append({
            'Test Name':line_data[0],
            'Test Status':line_data[1].lower().capitalize(),
            'Release': file_path.parts[base_folder_idx - 1],
            'V&V Test Report': v_v,
            'RC': file_path.parts[base_folder_idx + 1],
            'Owner': file_path.parts[base_folder_idx + 2],
            'File Path': str(file_path)
        })

    if return_df:
        return pd.DataFrame(data)
    else:
        return data

def read_rally_output(file_path, return_df=True):
    """Reads in an .csv that's outputted directly from a Rally query.

    Args:
        file_path (String): path to the .csv file
        return_df (bool, optional): If true, returns pandas dataframe. Else dict. Defaults to True.

    Returns:
        pd.DataFrame or dict: Test names and corresponding statuses
    """
    # As it's unformatted, can just load using existing pandas method
    df = pd.read_csv(file_path) 
    if return_df:
        return df
    else:
        return df.to_dict('records')
    

def read_rx_tests(folder_path, return_df=True):
    """Reads in all the Rx automatic test .txt files.
    
    Will recursively go through the folders and read in .txt files.
    Make sure everything is unzipped first; in future can add functionality to
    unzip automatically if needed.

    Args:
        folder_path (String or pathlib.Path): path to the main Rx folder.
        return_df (bool, optional): If true, returns pandas dataframe. Else dict. Defaults to True.

    Returns:
        pd.DataFrame or dict: Test names and corresponding statuses
    """
    print(f"Loading Rx test files from {folder_path}")
    if isinstance(folder_path, str):
        folder_path = Path(folder_path)

    # Read all .txt files recursively in the folders
    file_list = [file_name for file_name in folder_path.rglob("*.txt") if "__MACOSX" not in str(file_name)]

    # Load in the data from each file, add to a single list
    data = []
    for file_name in file_list:
        data.extend(_read_rx_txt(file_name, return_df=False))
    
    # make a dataframe for outputting
    df = pd.DataFrame(data)

    # Output as either pandas dataframe or dict, depending on return_df setting.
    if return_df:
        return df
    else:
        return df.to_dict('records')
    
def _read_rx_txt(file_path, return_df=True):
    """Reads in a single Rx automatic test .txt file.

    Mostly just for the `read_rx_tests()` method. Usually won't need to call yourself

    Args:
        file_path (String): path to the .txt file
        return_df (bool, optional): If true, returns pandas dataframe. Else dict. Defaults to True.

    Returns:
        pd.DataFrame or dict: Test names and corresponding statuses
    """
    # If receive str, convert to pathlib object
    if isinstance(file_path, str):
        file_path = Path(file_path)

    # Read in file using pathlib method, split on new line symbols
    lines = file_path.read_text().split("\n")

    # Remove blank lines
    lines = [line for line in lines if len(line) > 0]

    data = []
    for line in lines:
        if len(line) < 2:
            continue 
        # split on the | char; should only be one of them
        line_data = line.split("|")

        # For some reason if theres multiple of that char, only consider the last one
        if len(line_data) > 2:
            print(f"found line with multiple '|' chars: {line}")
            line_data = "".join(line_data[:-1]) + [line_data[-1]]
        
        # Get relevant data for other columns
        base_folder_idx = [i for i, s in enumerate(file_path.parts) if "Rx" in str(s)][0]
        
        # Also add the V&V Test Report info by extracting it from file name
        er_folder_name = file_path.parts[base_folder_idx - 2]
        v_v = re.findall("ER([0-9]+ v[0-9]+|[0-9]+v[0-9]+)", er_folder_name)[0].replace("ER", "")
        
        # If there's no space in the file name (like ER2228014v53), add one
        v_idx = v_v.find("v")
        if v_v[v_idx - 1] != " ":
            v_v = v_v[:v_idx] + " " + v_v[v_idx:]
        
        # Store data in list
        data.append({
            'Test Name':line_data[0],
            'Test Status':line_data[1].lower().capitalize(),
            'Release': file_path.parts[base_folder_idx - 1],
            'V&V Test Report': v_v,
            'RC': file_path.parts[base_folder_idx + 1],
            'Owner': file_path.parts[base_folder_idx + 2],
            'File Path': str(file_path)
        })

    if return_df:
        return pd.DataFrame(data)
    else:
        return data
    
    
def write_error_log(output_path, invalid_dfs):
    """Outputs dataframe with all invalid rows to csv file

    Args:
        output_path (String): file path of output csv
        invalid_dfs (list): list of dataframes that have invalid rows

    Output:
        Outputs error_df to a csv file
    """
    error_df = pd.DataFrame()
    
    for df in invalid_dfs:
        error_df = error_df.append(df)
    
    error_df.to_csv(output_path, index=False)   
    
    
    
